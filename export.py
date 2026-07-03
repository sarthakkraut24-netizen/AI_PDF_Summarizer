# ==========================================================
# EXPORT.PY
# AI Document Assistant
# Compatible with Python 3.13+
# ==========================================================

import os

from reportlab.lib.pagesizes import letter

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from reportlab.lib.styles import getSampleStyleSheet

from docx import Document


# ==========================================================
# CREATE OUTPUT DIRECTORY
# ==========================================================

OUTPUT_FOLDER = "exports"

os.makedirs(

    OUTPUT_FOLDER,

    exist_ok=True

)


# ==========================================================
# EXPORT TO TXT
# ==========================================================

def export_txt(

    text,

    filename="summary.txt"

):

    path = os.path.join(

        OUTPUT_FOLDER,

        filename

    )

    with open(

        path,

        "w",

        encoding="utf-8"

    ) as file:

        file.write(text)

    return path


# ==========================================================
# EXPORT TO DOCX
# ==========================================================

def export_docx(

    text,

    filename="summary.docx"

):

    document = Document()

    document.add_heading(

        "AI Document Assistant",

        level=1

    )

    document.add_paragraph(text)

    path = os.path.join(

        OUTPUT_FOLDER,

        filename

    )

    document.save(path)

    return path


# ==========================================================
# EXPORT TO PDF
# ==========================================================

def export_pdf(

    text,

    filename="summary.pdf"

):

    path = os.path.join(

        OUTPUT_FOLDER,

        filename

    )

    document = SimpleDocTemplate(

        path,

        pagesize=letter

    )

    styles = getSampleStyleSheet()

    story = []

    story.append(

        Paragraph(

            "<b>AI Document Assistant</b>",

            styles["Heading1"]

        )

    )

    paragraphs = text.split("\n")

    for para in paragraphs:

        if para.strip():

            story.append(

                Paragraph(

                    para,

                    styles["BodyText"]

                )

            )

    document.build(story)

    return path


# ==========================================================
# AUTO EXPORT
# ==========================================================

def export_file(

    text,

    file_type="PDF",

    filename="summary"

):

    file_type = file_type.lower()

    if file_type == "pdf":

        return export_pdf(

            text,

            filename + ".pdf"

        )

    elif file_type == "docx":

        return export_docx(

            text,

            filename + ".docx"

        )

    elif file_type == "txt":

        return export_txt(

            text,

            filename + ".txt"

        )

    else:

        raise ValueError(

            "Supported formats: PDF, DOCX, TXT"

        )


# ==========================================================
# TEST
# ==========================================================

if __name__ == "__main__":

    sample = """

    Artificial Intelligence is changing
    healthcare, education and finance.

    """

    print(

        export_pdf(

            sample

        )

    )

    print(

        export_docx(

            sample

        )

    )

    print(

        export_txt(

            sample

        )

    )